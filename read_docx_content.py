#!/usr/bin/env python3
# -*- coding: utf-8 -*-
try:
    from docx import Document
except ImportError:
    try:
        import docx
        Document = docx.Document
    except ImportError:
        print("ERROR: python-docx library is not installed.")
        print("Please install it with: pip install python-docx")
        exit(1)

import sys
import os

def read_docx(file_path):
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None
            
        doc = Document(file_path)
        full_text = []
        
        # Читаем параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Читаем таблицы
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        text = '\n'.join(full_text)
        return text
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        import traceback
        traceback.print_exc()
        return None

if __name__ == "__main__":
    file_path = "Сайт_ТГИД.docx"
    if not os.path.exists(file_path):
        # Пробуем найти файл в текущей директории
        current_dir = os.getcwd()
        print(f"Current directory: {current_dir}")
        print(f"Looking for file: {file_path}")
        
        # Список файлов в текущей директории
        files = [f for f in os.listdir('.') if f.endswith('.docx')]
        print(f"Found .docx files: {files}")
        
        if files:
            file_path = files[0]
            print(f"Using file: {file_path}")
    
    content = read_docx(file_path)
    if content:
        print("\n" + "="*80)
        print("DOCUMENT CONTENT:")
        print("="*80 + "\n")
        print(content)
        print("\n" + "="*80)
    else:
        print("Failed to read document")

